import os
import telebot
from telebot import types
import requests
import pandas as pd
import docx

bot_token = '6975517932:AAFwFndSlKz5MCNf32t12LwCTDDFT4VtCc8'
bot = telebot.TeleBot(bot_token)
api_key = 'AIzaSyA0rmJ8oxfyylSYZjtspyO2a7ggVjc41Dg'
reading_list = []

@bot.message_handler(commands=['start'])
def start(message):
    bot.send_message(message.chat.id, "Welcome to PagePal! Type /help to see the list of commands.")

@bot.message_handler(commands=['help'])
def help(message):
    help_text = (
        "/start - Welcome message\n"
        "/book - Search for books by genre (returns CSV format)\n"
        "/preview - Get a preview link for a book\n"
        "/list - Add or remove a book from the reading list\n"
        "/reading_list - Show your reading list (returns DOCX format)"
    )
    bot.send_message(message.chat.id, help_text)

@bot.message_handler(commands=['preview'])
def preview(message):
    bot.send_message(message.chat.id, "Enter the title of the book you want to preview:")
    bot.register_next_step_handler(message, get_preview)

def get_preview(message):
    title = message.text
    params = {"q": f"intitle:{title}", "key": "AIzaSyA0rmJ8oxfyylSYZjtspyO2a7ggVjc41Dg"} 
    response = requests.get("https://www.googleapis.com/books/v1/volumes", params=params)
    data = response.json()
    
    if "items" in data:
        volume_info = data["items"][0]["volumeInfo"]
        preview_link = volume_info.get("previewLink", "No preview available")
        bot.send_message(message.chat.id, f"Preview: {preview_link}")
    else:
        bot.send_message(message.chat.id, "Sorry, we don't have a preview for that book.")

@bot.message_handler(commands=['book'])
def book(message):
    bot.send_message(message.chat.id, "Enter the genre of the book you're looking for (e.g., fantasy, science fiction):")
    bot.register_next_step_handler(message, search_books)

def search_books(message):
    genre = message.text

    params = {'q': f'subject:{genre}', 'key': "AIzaSyA0rmJ8oxfyylSYZjtspyO2a7ggVjc41Dg", 'maxResults': 5}
    response = requests.get("https://www.googleapis.com/books/v1/volumes", params=params)
    data = response.json()
    books = []
    
    for item in data.get('items', []):
        book_info = {
            'Title': item['volumeInfo'].get('title', 'N/A'),
            'Authors': ', '.join(item['volumeInfo'].get('authors', 'N/A')),
            'Publisher': item['volumeInfo'].get('publisher', 'N/A'),
            'PublishedDate': item['volumeInfo'].get('publishedDate', 'N/A'),
            'Preview Link': item['volumeInfo'].get('previewLink', 'N/A')
        }
        books.append(book_info)
    
    if books:
        df = pd.DataFrame(books)
        csv_file = 'book_recommendations.csv'
        df.to_csv(csv_file, index=False)
        with open(csv_file, 'rb') as f:
            bot.send_document(message.chat.id, f)
        os.remove(csv_file)
    else:
        bot.send_message(message.chat.id, "No books found for the specified genre.")

@bot.message_handler(commands=['list'])
def modify_list(message):
    bot.send_message(message.chat.id, "Enter the title of the book to add or remove from your reading list:")
    bot.register_next_step_handler(message, update_reading_list)

def update_reading_list(message):
    book_title = message.text
    if book_title in reading_list:
        reading_list.remove(book_title)
        bot.send_message(message.chat.id, f"Removed '{book_title}' from your reading list.")
    else:
        reading_list.append(book_title)
        bot.send_message(message.chat.id, f"Added '{book_title}' to your reading list.")

@bot.message_handler(commands=['reading_list'])
def show_reading_list(message):
    if reading_list:
        doc = docx.Document()
        doc.add_heading('Your Reading List', 0)
        for book in reading_list:
            doc.add_paragraph(book)
        docx_file = 'reading_list.docx'
        doc.save(docx_file)
        with open(docx_file, 'rb') as f:
            bot.send_document(message.chat.id, f)
        os.remove(docx_file)
    else:
        bot.send_message(message.chat.id, "Your reading list is empty.")

@bot.message_handler(func=lambda message:True)
def wrong_msg(message):
    if message.text.startswith("/"):
        bot.send_message(message.chat.id, "I don't know that command. Plz enter a valid command")
    else:
        bot.send_message(message.chat.id, "All commmands must start with /")
bot.polling()
